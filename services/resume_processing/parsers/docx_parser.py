"""
Парсер текста резюме DOCX.

# Русский комментарий:
Этот модуль предоставляет функции для извлечения текста из DOCX (Microsoft Word) файлов резюме
с использованием библиотеки python-docx.
"""
import logging
from pathlib import Path
from typing import Dict, Optional, Union

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: Union[str, Path]) -> Dict[str, Optional[str]]:
    """
    Извлечь текст из DOCX файла Microsoft Word.

    Extract text from a Microsoft Word DOCX file.

    Args:
        file_path: Путь к DOCX файлу / Path to the DOCX file

    Returns:
        Словарь содержащий:
            - text: Извлеченный текст (None при ошибке) / Extracted text content (None if extraction fails)
            - method: Всегда 'python-docx' / Always 'python-docx'
            - paragraphs: Количество параграфов / Number of paragraphs detected
            - tables: Количество таблиц / Number of tables detected
            - error: Сообщение об ошибке при неудаче / Error message if extraction fails

    Raises:
        FileNotFoundError: Если файл не существует / If the file doesn't exist
        ValueError: Если файл не является валидным DOCX / If the file is not a valid DOCX

    Examples:
        >>> result = extract_text_from_docx("resume.docx")
        >>> print(result["text"])
        'John Doe\\nSoftware Engineer...'
        >>> print(result["paragraphs"])
        25
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not file_path.suffix.lower() in (".docx", ".doc"):
        raise ValueError(f"File is not a DOCX: {file_path}")

    try:
        doc = Document(str(file_path))

        # Извлечение текста из параграфов / Extract text from paragraphs
        paragraph_texts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph_texts.append(paragraph.text)

        # Извлечение текста из таблиц / Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        # Объединение текста из параграфов и таблиц / Combine text from paragraphs and tables
        all_texts = paragraph_texts + table_texts
        text = "\n\n".join(all_texts)

        logger.info(
            f"Extracted {len(text)} chars from {file_path.name} "
            f"({len(paragraph_texts)} paragraphs, {len(doc.tables)} tables)"
        )

        return {
            "text": text if text.strip() else None,
            "method": "python-docx",
            "paragraphs": len(paragraph_texts),
            "tables": len(doc.tables),
            "error": None,
        }

    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return {
            "text": None,
            "method": None,
            "paragraphs": 0,
            "tables": 0,
            "error": f"DOCX extraction failed: {str(e)}",
        }


def extract_text_with_metadata(file_path: Union[str, Path]) -> Dict[str, Optional[Union[str, int]]]:
    """
    Извлечь текст и метаданные из DOCX файла.

    Extract text and metadata from a DOCX file.

    Эта функция дополнительно извлекает метаданные документа такие как
    автор, заголовок, и дата создания если они доступны.

    This function additionally extracts document metadata such as author,
    title, and creation date if available.

    Args:
        file_path: Путь к DOCX файлу / Path to the DOCX file

    Returns:
        Расширенный словарь содержащий:
            - text: Извлеченный текст / Extracted text content
            - method: Метод извлечения / Extraction method
            - paragraphs: Количество параграфов / Number of paragraphs
            - tables: Количество таблиц / Number of tables
            - author: Автор документа (если доступен) / Document author if available
            - title: Заголовок документа (если доступен) / Document title if available
            - created: Дата создания (если доступна) / Creation date if available
            - error: Сообщение об ошибке / Error message

    Examples:
        >>> result = extract_text_with_metadata("resume.docx")
        >>> print(result["author"])
        'John Doe'
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        doc = Document(str(file_path))

        # Извлечение текста / Extract text
        base_result = extract_text_from_docx(file_path)

        # Извлечение метаданных ядра документа / Extract core document metadata
        core_props = doc.core_properties

        result = {
            **base_result,
            "author": core_props.author or None,
            "title": core_props.title or None,
            "subject": core_props.subject or None,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "last_modified_by": core_props.last_modified_by or None,
        }

        logger.info(
            f"Extracted metadata from {file_path.name}: "
            f"author={result['author']}, title={result['title']}"
        )

        return result

    except Exception as e:
        logger.error(f"Error extracting metadata from DOCX {file_path}: {e}")
        return {
            "text": None,
            "method": None,
            "paragraphs": 0,
            "tables": 0,
            "author": None,
            "title": None,
            "subject": None,
            "created": None,
            "modified": None,
            "last_modified_by": None,
            "error": f"DOCX metadata extraction failed: {str(e)}",
        }


def validate_docx_file(file_path: Union[str, Path]) -> Dict[str, Union[bool, str]]:
    """
    Валидировать DOCX файл перед извлечением текста.

    Validate a DOCX file before extraction.

    Проверки:
        - Файл существует / File exists
        - Имеет расширение .docx или .doc / Has .docx or .doc extension
        - Не пустой / Is not empty
        - Может быть открыт библиотекой python-docx / Can be opened by python-docx

    Args:
        file_path: Путь к DOCX файлу / Path to the DOCX file

    Returns:
        Словарь с результатами валидации:
            - valid: Булево значение указывающее валидность файла / Boolean indicating if file is valid
            - reason: Строка объясняющая почему валидация не прошла (если применимо) / String explaining why validation failed

    Examples:
        >>> validation = validate_docx_file("resume.docx")
        >>> if validation["valid"]:
        ...     print("File is valid")
    """
    file_path = Path(file_path)

    # Проверка существования файла / Check file exists
    if not file_path.exists():
        return {
            "valid": False,
            "reason": f"File does not exist: {file_path}"
        }

    # Проверка расширения / Check extension
    if file_path.suffix.lower() not in (".docx", ".doc"):
        return {
            "valid": False,
            "reason": f"File does not have .docx or .doc extension: {file_path.suffix}"
        }

    # Проверка размера файла / Check file size
    if file_path.stat().st_size == 0:
        return {
            "valid": False,
            "reason": "File is empty (0 bytes)"
        }

    # Проверка что файл может быть открыт / Check file can be opened
    try:
        doc = Document(str(file_path))
        # Проверка что есть хотя бы один элемент / Check there's at least one element
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            return {
                "valid": False,
                "reason": "DOCX contains no paragraphs or tables"
            }
    except Exception as e:
        return {
            "valid": False,
            "reason": f"Cannot open DOCX file: {str(e)}"
        }

    return {"valid": True, "reason": None}
