"""
Tests for data extraction service.

Tests PDF and DOCX text extraction with various edge cases including:
- Valid files
- Malformed files
- Empty files
- Missing files
- Files with special characters
- Files with tables (DOCX)
"""
import os
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch

import pytest
from PyPDF2 import PdfReader
from docx import Document

# Import from parent directory
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from extract import (
    extract_text_from_pdf,
    extract_text_from_docx,
    validate_pdf_file,
    validate_docx_file,
    _extract_with_pypdf2,
    _extract_with_pdfplumber,
    _extract_with_python_docx,
)


class TestPDFExtraction:
    """Test PDF text extraction functionality."""

    def test_extract_text_from_valid_pdf(self, sample_pdf_path):
        """Test extraction from a valid PDF file."""
        result = extract_text_from_pdf(sample_pdf_path)

        assert result["text"] is not None
        assert len(result["text"]) > 0
        assert result["method"] in ["pypdf2", "pdfplumber"]
        assert result["pages"] > 0
        assert result["error"] is None

    def test_extract_text_from_pdf_with_fallback(self, sample_pdf_path):
        """Test PDF extraction with pdfplumber fallback."""
        result = extract_text_from_pdf(sample_pdf_path, use_fallback=True)

        assert result["text"] is not None
        assert result["method"] in ["pypdf2", "pdfplumber"]

    def test_extract_text_from_pdf_without_fallback(self, sample_pdf_path):
        """Test PDF extraction without fallback."""
        result = extract_text_from_pdf(sample_pdf_path, use_fallback=False)

        assert result["text"] is not None
        # Should use PyPDF2 only
        assert result["method"] == "pypdf2"

    def test_extract_text_from_nonexistent_pdf(self):
        """Test extraction from non-existent file raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            extract_text_from_pdf("nonexistent.pdf")

    def test_extract_text_from_invalid_extension(self, sample_pdf_path, tmp_path):
        """Test extraction from file with wrong extension raises ValueError."""
        # Copy PDF but rename to .txt
        invalid_file = tmp_path / "document.txt"
        if sample_pdf_path.exists():
            import shutil
            shutil.copy(sample_pdf_path, invalid_file)

        with pytest.raises(ValueError, match="not a PDF"):
            extract_text_from_pdf(invalid_file)

    def test_extract_text_from_empty_pdf(self, empty_pdf_path):
        """Test extraction from empty PDF returns appropriate error."""
        result = extract_text_from_pdf(empty_pdf_path)

        assert result["text"] is None
        assert result["error"] is not None

    def test_extract_with_corrupted_pdf(self, tmp_path):
        """Test extraction from corrupted PDF returns error gracefully."""
        corrupted_file = tmp_path / "corrupted.pdf"
        corrupted_file.write_bytes(b"This is not a valid PDF file")

        result = extract_text_from_pdf(corrupted_file)

        assert result["text"] is None
        assert result["method"] is None
        assert result["error"] is not None

    def test_validate_valid_pdf(self, sample_pdf_path):
        """Test validation of valid PDF file."""
        result = validate_pdf_file(sample_pdf_path)

        assert result["valid"] is True
        assert result["reason"] is None

    def test_validate_nonexistent_pdf(self):
        """Test validation of non-existent PDF file."""
        result = validate_pdf_file("nonexistent.pdf")

        assert result["valid"] is False
        assert result["reason"] == "File not found"

    def test_validate_wrong_extension(self, tmp_path):
        """Test validation of file with wrong extension."""
        wrong_file = tmp_path / "document.txt"
        wrong_file.write_text("Not a PDF")

        result = validate_pdf_file(wrong_file)

        assert result["valid"] is False
        assert result["reason"] == "Not a PDF file"

    def test_validate_empty_pdf(self, empty_pdf_path):
        """Test validation of empty PDF file."""
        result = validate_pdf_file(empty_pdf_path)

        assert result["valid"] is False
        assert result["reason"] in ["File is empty", "PDF has no pages"]


class TestDOCXExtraction:
    """Test DOCX text extraction functionality."""

    def test_extract_text_from_valid_docx(self, sample_docx_path):
        """Test extraction from a valid DOCX file."""
        result = extract_text_from_docx(sample_docx_path)

        assert result["text"] is not None
        assert len(result["text"]) > 0
        assert result["method"] == "python-docx"
        assert result["paragraphs"] > 0
        assert result["error"] is None

    def test_extract_text_from_docx_with_tables(self, sample_docx_with_tables_path):
        """Test extraction from DOCX file with tables."""
        result = extract_text_from_docx(sample_docx_with_tables_path)

        assert result["text"] is not None
        assert len(result["text"]) > 0
        # Tables should be extracted and included in text
        assert result["method"] == "python-docx"

    def test_extract_text_from_nonexistent_docx(self):
        """Test extraction from non-existent DOCX raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            extract_text_from_docx("nonexistent.docx")

    def test_extract_text_from_invalid_docx_extension(self, tmp_path):
        """Test extraction from file with wrong extension raises ValueError."""
        wrong_file = tmp_path / "document.txt"
        wrong_file.write_text("Not a DOCX")

        with pytest.raises(ValueError, match="not a DOCX"):
            extract_text_from_docx(wrong_file)

    def test_extract_text_from_empty_docx(self, empty_docx_path):
        """Test extraction from empty DOCX file."""
        result = extract_text_from_docx(empty_docx_path)

        # Empty DOCX should return empty text or None
        assert result["text"] is None or len(result["text"].strip()) == 0

    def test_extract_text_from_corrupted_docx(self, tmp_path):
        """Test extraction from corrupted DOCX returns error gracefully."""
        corrupted_file = tmp_path / "corrupted.docx"
        corrupted_file.write_bytes(b"This is not a valid DOCX file")

        result = extract_text_from_docx(corrupted_file)

        assert result["text"] is None
        assert result["method"] is None
        assert result["error"] is not None

    def test_validate_valid_docx(self, sample_docx_path):
        """Test validation of valid DOCX file."""
        result = validate_docx_file(sample_docx_path)

        assert result["valid"] is True
        assert result["reason"] is None

    def test_validate_nonexistent_docx(self):
        """Test validation of non-existent DOCX file."""
        result = validate_docx_file("nonexistent.docx")

        assert result["valid"] is False
        assert result["reason"] == "File not found"

    def test_validate_wrong_docx_extension(self, tmp_path):
        """Test validation of file with wrong extension."""
        wrong_file = tmp_path / "document.txt"
        wrong_file.write_text("Not a DOCX")

        result = validate_docx_file(wrong_file)

        assert result["valid"] is False
        assert result["reason"] == "Not a DOCX file"

    def test_validate_empty_docx(self, empty_docx_path):
        """Test validation of empty DOCX file."""
        result = validate_docx_file(empty_docx_path)

        assert result["valid"] is False
        assert result["reason"] in ["File is empty", "DOCX has no content"]


class TestErrorHandling:
    """Test error handling for malformed and edge case files."""

    def test_docx_with_special_characters(self, tmp_path):
        """Test DOCX extraction with special unicode characters."""
        doc = Document()
        doc.add_paragraph("Тест на русском")  # Russian
        doc.add_paragraph("Test in English")
        doc.add_paragraph("中文测试")  # Chinese

        docx_file = tmp_path / "special_chars.docx"
        doc.save(str(docx_file))

        result = extract_text_from_docx(docx_file)

        # Should extract text with special characters
        assert result["text"] is not None
        assert len(result["text"]) > 0
        assert result["error"] is None

    def test_pdf_with_no_extractable_text(self, tmp_path):
        """Test PDF with minimal/extractable text."""
        # Create a minimal PDF with very little text
        pdf_file = tmp_path / "minimal.pdf"
        pdf_file.write_bytes(
            b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 << /Type /Font "
            b"/Subtype /Type1 /BaseFont /Helvetica >> >> >> >>\nendobj\n"
            b"4 0 obj\n<< /Length 44 >>\nstream\n"
            b"BT\n/F1 12 Tf\n100 700 Td\n(X) Tj\nET\n"
            b"endstream\nendobj\n"
            b"xref\n0 5\n0000000000 65535 f\n0000000009 00000 n\n0000000058 00000 n\n"
            b"0000000115 00000 n\n0000000300 00000 n\ntrailer\n<< /Size 5 /Root 1 0 R >>\n"
            b"startxref\n403\n%%EOF\n"
        )

        result = extract_text_from_pdf(pdf_file)

        # Should handle gracefully - might return empty text or very short text
        assert result["error"] is None  # No crash, just empty result

    def test_docx_with_only_tables(self, tmp_path):
        """Test DOCX file that only contains tables (common in resume templates)."""
        doc = Document()

        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "John Doe"
        table.rows[1].cells[0].text = "Email"
        table.rows[1].cells[1].text = "john@example.com"
        table.rows[2].cells[0].text = "Phone"
        table.rows[2].cells[1].text = "555-1234"

        docx_file = tmp_path / "only_tables.docx"
        doc.save(str(docx_file))

        result = extract_text_from_docx(docx_file)

        # Should extract table contents
        assert result["text"] is not None
        assert "Name" in result["text"] or "John" in result["text"]


# ===== Fixtures =====

@pytest.fixture
def sample_pdf_path():
    """Use the pre-created sample PDF file from test_samples."""
    sample_file = Path(__file__).parent.parent / "test_samples" / "sample.pdf"
    if not sample_file.exists():
        pytest.skip(f"Sample PDF not found at {sample_file}")
    return sample_file


@pytest.fixture
def empty_pdf_path(tmp_path):
    """Create an empty PDF file for testing."""
    # Create a minimal empty PDF
    pdf_file = tmp_path / "empty.pdf"
    pdf_file.write_bytes(
        b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 << /Type /Font /Subtype /Type1 "
        b"/BaseFont /Helvetica >> >> >> >>\nendobj\nxref\n0 4\n"
        b"0000000000 65535 f\n0000000009 00000 n\n0000000058 00000 n\n"
        b"0000000115 00000 n\ntrailer\n<< /Size 4 /Root 1 0 R >>\n"
        b"startxref\n210\n%%EOF\n"
    )
    return pdf_file


@pytest.fixture
def sample_docx_path():
    """Use the pre-created sample DOCX file from test_samples."""
    sample_file = Path(__file__).parent.parent / "test_samples" / "sample.docx"
    if not sample_file.exists():
        pytest.skip(f"Sample DOCX not found at {sample_file}")
    return sample_file


@pytest.fixture
def sample_docx_with_tables_path(tmp_path):
    """Create a DOCX file with tables for testing."""
    doc = Document()
    doc.add_paragraph("Resume")

    # Add a table
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Jane Smith"
    table.rows[1].cells[0].text = "Position"
    table.rows[1].cells[1].text = "Data Scientist"
    table.rows[2].cells[0].text = "Email"
    table.rows[2].cells[1].text = "jane@example.com"
    table.rows[3].cells[0].text = "Phone"
    table.rows[3].cells[1].text = "555-5678"

    docx_file = tmp_path / "with_tables.docx"
    doc.save(str(docx_file))

    return docx_file


@pytest.fixture
def empty_docx_path(tmp_path):
    """Create an empty DOCX file for testing."""
    doc = Document()
    # Don't add any content

    docx_file = tmp_path / "empty.docx"
    doc.save(str(docx_file))

    return docx_file
