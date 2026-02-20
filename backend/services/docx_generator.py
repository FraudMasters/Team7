"""
DOCX Generator Service for Resume Templates.

This module provides functionality to generate Word documents (DOCX)
from resume data, supporting professional formatting and ATS-friendly output.

The service supports:
- DOCX generation using python-docx library
- Professional resume formatting with proper styling
- Multiple page formats (A4, Letter)
- Custom margins and document settings
- ATS-friendly document structure
- Document metadata (author, title, etc.)
- Error handling and graceful fallback

The service takes resume content data and creates well-formatted
Word documents that can be downloaded by job seekers.

Example:
    >>> from services.docx_generator import DOCXGenerator
    >>> from schemas.resume_builder import ResumeContent
    >>> generator = DOCXGenerator()
    >>> result = await generator.generate_resume_docx(
    ...     resume_content=resume_content,
    ...     candidate_name="John Doe"
    ... )
    >>> if result.success:
    ...     with open("resume.docx", "wb") as f:
    ...         f.write(result.docx_bytes)
"""
import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Dict, List, Optional

from config import get_settings

logger = logging.getLogger(__name__)

# Global DOCX generator instance
_docx_generator: Optional["DOCXGenerator"] = None


@dataclass
class DOCXGenerationOptions:
    """
    Configuration options for DOCX generation.

    Attributes:
        page_format: Page size format ('A4' or 'Letter')
        margin_top: Top margin in points
        margin_bottom: Bottom margin in points
        margin_left: Left margin in points
        margin_right: Right margin in points
        font_family: Primary font family
        font_size_body: Font size for body text in points
        font_size_heading: Font size for headings in points
        font_size_name: Font size for candidate name in points
        include_page_numbers: Whether to include page numbers
        line_spacing: Line spacing factor (1.0 = single, 1.5 = 1.5 lines)
        accent_color: Primary accent color for headings (hex string)
    """
    page_format: str = "A4"
    margin_top: float = 36.0  # 0.5 inch in points
    margin_bottom: float = 36.0
    margin_left: float = 54.0  # 0.75 inch in points
    margin_right: float = 54.0
    font_family: str = "Calibri"
    font_size_body: float = 11.0
    font_size_heading: float = 14.0
    font_size_name: float = 20.0
    include_page_numbers: bool = False
    line_spacing: float = 1.15
    accent_color: str = "2E5090"  # Professional blue

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "page_format": self.page_format,
            "margin_top": self.margin_top,
            "margin_bottom": self.margin_bottom,
            "margin_left": self.margin_left,
            "margin_right": self.margin_right,
            "font_family": self.font_family,
            "font_size_body": self.font_size_body,
            "font_size_heading": self.font_size_heading,
            "font_size_name": self.font_size_name,
            "include_page_numbers": self.include_page_numbers,
            "line_spacing": self.line_spacing,
            "accent_color": self.accent_color,
        }


@dataclass
class DOCXMetadata:
    """
    Metadata for generated DOCX documents.

    Attributes:
        title: Document title
        author: Document author
        subject: Document subject/description
        keywords: List of keywords for the document
        creator: Software that created the document
        creation_date: When the document was created
    """
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    creator: str = "AgentHR Resume Builder"
    creation_date: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        metadata = {}
        if self.title:
            metadata["title"] = self.title
        if self.author:
            metadata["author"] = self.author
        if self.subject:
            metadata["subject"] = self.subject
        if self.keywords:
            metadata["keywords"] = ", ".join(self.keywords)
        if self.creator:
            metadata["creator"] = self.creator
        if self.creation_date:
            metadata["created"] = self.creation_date
        return metadata


@dataclass
class DOCXGenerationResult:
    """
    Result of DOCX generation.

    Attributes:
        success: Whether DOCX generation succeeded
        docx_bytes: Generated DOCX content as bytes
        filename: Suggested filename for the DOCX
        content_type: MIME type (application/vnd.openxmlformats-officedocument.wordprocessingml.document)
        error_message: Error message if generation failed
        file_size: Size of generated DOCX in bytes
        page_count: Estimated number of pages
        metadata: Document metadata included in the document
    """
    success: bool
    docx_bytes: Optional[bytes] = None
    filename: Optional[str] = None
    content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    error_message: Optional[str] = None
    file_size: int = 0
    page_count: int = 0
    metadata: Optional[DOCXMetadata] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "success": self.success,
            "filename": self.filename,
            "content_type": self.content_type,
            "error_message": self.error_message,
            "file_size": self.file_size,
            "page_count": self.page_count,
            "metadata": self.metadata.to_dict() if self.metadata else None,
        }


class DOCXGenerator:
    """
    DOCX Generator Service for Resume Templates.

    This service converts resume data into professional Word documents
    suitable for job applications.

    The service uses python-docx for document generation with support for:
    - Professional formatting and styling
    - Proper section structure for ATS compatibility
    - Custom fonts and sizes
    - Document metadata (title, author, etc.)
    - ATS-friendly output (clean, simple structure)

    python-docx is used because:
    - Native DOCX format support
    - Good control over document structure
    - Wide compatibility with Word processors
    - ATS-friendly document format

    Attributes:
        enabled: Whether DOCX generation is enabled
        default_options: Default DOCX generation options
        docx_available: Whether python-docx library is available

    Example:
        >>> generator = DOCXGenerator()
        >>> result = await generator.generate_resume_docx(
        ...     resume_content=resume_content,
        ...     filename="John_Doe_Resume.docx"
        ... )
        >>> if result.success:
        ...     with open("resume.docx", "wb") as f:
        ...         f.write(result.docx_bytes)
    """

    # Page format constants
    PAGE_FORMAT_A4 = "A4"
    PAGE_FORMAT_LETTER = "Letter"

    # Page sizes in inches (for python-docx)
    PAGE_SIZES = {
        "A4": (8.27, 11.69),  # Width, Height in inches
        "Letter": (8.5, 11.0),
    }

    def __init__(
        self,
        enabled: Optional[bool] = None,
        default_options: Optional[DOCXGenerationOptions] = None,
    ) -> None:
        """
        Initialize the DOCX generator service.

        Args:
            enabled: Whether DOCX generation is enabled
            default_options: Default DOCX generation options
        """
        settings = get_settings()

        self.enabled = enabled if enabled is not None else True
        self.default_options = default_options or DOCXGenerationOptions()

        # Try to import python-docx, disable if not available
        self.docx_available = False
        if self.enabled:
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.style import WD_STYLE_TYPE

                self._Document = Document
                self._Inches = Inches
                self._Pt = Pt
                self._RGBColor = RGBColor
                self._WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
                self._WD_STYLE_TYPE = WD_STYLE_TYPE
                self.docx_available = True
                logger.info("DOCXGenerator initialized with python-docx")
            except ImportError:
                logger.warning(
                    "python-docx not installed, DOCX generation will be disabled. "
                    "Install with: pip install python-docx"
                )
                self.docx_available = False
        else:
            logger.info("DOCXGenerator initialized but disabled")

    def is_available(self) -> bool:
        """
        Check if DOCX generation is available.

        Returns:
            True if python-docx is available and enabled
        """
        return self.enabled and self.docx_available

    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename by removing/replacing problematic characters.

        Args:
            filename: Original filename

        Returns:
            Sanitized filename safe for filesystem use
        """
        # Replace problematic characters with underscore
        problematic = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        sanitized = filename
        for char in problematic:
            sanitized = sanitized.replace(char, '_')

        # Ensure filename ends with .docx
        if not sanitized.lower().endswith('.docx'):
            sanitized += '.docx'

        return sanitized

    def _generate_suggested_filename(
        self,
        candidate_name: Optional[str] = None,
        template_name: Optional[str] = None,
    ) -> str:
        """
        Generate a suggested filename for the resume DOCX.

        Args:
            candidate_name: Name of the candidate
            template_name: Name of the template used

        Returns:
            Suggested filename in format: "FirstName_LastName_Resume_Template.docx"
        """
        parts = []

        if candidate_name:
            # Convert name to filename-friendly format
            name_parts = candidate_name.strip().split()
            if name_parts:
                formatted_name = "_".join(
                    part.capitalize() for part in name_parts if part
                )
                parts.append(formatted_name)

        if not parts:
            parts.append("Resume")

        if template_name:
            template_formatted = template_name.strip().replace(" ", "_")
            parts.append(template_formatted)

        if len(parts) == 1:
            parts.append("Resume")

        filename = "_".join(parts)
        return self._sanitize_filename(filename)

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """
        Convert hex color string to RGB tuple.

        Args:
            hex_color: Hex color string (with or without #)

        Returns:
            Tuple of (R, G, B) values
        """
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def _add_contact_line(
        self,
        paragraph,
        text: str,
        separator: str = " | ",
    ) -> None:
        """
        Add text to a contact info paragraph with separator.

        Args:
            paragraph: The paragraph to add text to
            text: Text to add
            separator: Separator to use before text
        """
        if text:
            run = paragraph.add_run(separator + text)
            run.font.size = self._Pt(10)

    async def generate_resume_docx(
        self,
        resume_content: Dict[str, Any],
        filename: Optional[str] = None,
        candidate_name: Optional[str] = None,
        template_name: Optional[str] = None,
        metadata: Optional[DOCXMetadata] = None,
        options: Optional[DOCXGenerationOptions] = None,
    ) -> DOCXGenerationResult:
        """
        Generate a DOCX resume from resume content data.

        Args:
            resume_content: Dictionary containing resume data (ResumeContent structure)
            filename: Optional filename for the DOCX (auto-generated if not provided)
            candidate_name: Name of the candidate for filename and metadata
            template_name: Name of the template used
            metadata: Optional DOCX metadata
            options: Optional DOCX generation options

        Returns:
            DOCXGenerationResult with generated DOCX or error details

        Example:
            >>> result = await generator.generate_resume_docx(
            ...     resume_content={
            ...         "personal_info": {"full_name": "John Doe", "email": "john@example.com"},
            ...         "work_experience": [...],
            ...         "education": [...],
            ...         "skills": [...]
            ...     },
            ...     candidate_name="John Doe"
            ... )
        """
        # Check if python-docx is available
        if not self.is_available():
            return DOCXGenerationResult(
                success=False,
                error_message="python-docx is not available. Install with: pip install python-docx"
            )

        # Validate resume content
        if not resume_content:
            return DOCXGenerationResult(
                success=False,
                error_message="Resume content is empty"
            )

        # Use provided options or defaults
        generation_options = options or self.default_options

        # Extract personal info for candidate name if not provided
        personal_info = resume_content.get("personal_info", {})
        if not candidate_name:
            candidate_name = personal_info.get("full_name")

        # Generate filename if not provided
        if not filename:
            filename = self._generate_suggested_filename(
                candidate_name=candidate_name,
                template_name=template_name
            )
        else:
            filename = self._sanitize_filename(filename)

        # Prepare metadata
        if not metadata:
            metadata = DOCXMetadata(
                title=f"{candidate_name} - Resume" if candidate_name else "Resume",
                author=candidate_name or "Candidate",
                subject="Professional Resume",
                keywords=["resume", "cv", "curriculum vitae"],
                creation_date=datetime.utcnow().isoformat(),
            )

        try:
            logger.info(
                f"Generating DOCX resume: {filename} "
                f"(page_format={generation_options.page_format})"
            )

            # Create document
            doc = self._Document()

            # Set document properties
            doc.core_properties.title = metadata.title or "Resume"
            doc.core_properties.author = metadata.author or "Candidate"
            doc.core_properties.subject = metadata.subject or "Professional Resume"
            doc.core_properties.creator = metadata.creator

            # Set page size
            page_width, page_height = self.PAGE_SIZES.get(
                generation_options.page_format,
                self.PAGE_SIZES["A4"]
            )
            for section in doc.sections:
                section.page_width = self._Inches(page_width)
                section.page_height = self._Inches(page_height)
                section.top_margin = self._Inches(generation_options.margin_top / 72)
                section.bottom_margin = self._Inches(generation_options.margin_bottom / 72)
                section.left_margin = self._Inches(generation_options.margin_left / 72)
                section.right_margin = self._Inches(generation_options.margin_right / 72)

            # Get accent color
            accent_rgb = self._hex_to_rgb(generation_options.accent_color)

            # Add header section with name and contact info
            self._add_header_section(doc, personal_info, generation_options, accent_rgb)

            # Add summary section
            summary = resume_content.get("summary") or personal_info.get("summary")
            if summary:
                self._add_summary_section(doc, summary, generation_options, accent_rgb)

            # Add work experience section
            work_experience = resume_content.get("work_experience", [])
            if work_experience:
                self._add_work_experience_section(doc, work_experience, generation_options, accent_rgb)

            # Add education section
            education = resume_content.get("education", [])
            if education:
                self._add_education_section(doc, education, generation_options, accent_rgb)

            # Add skills section
            skills = resume_content.get("skills", [])
            if skills:
                self._add_skills_section(doc, skills, generation_options, accent_rgb)

            # Add certifications section
            certifications = resume_content.get("certifications", [])
            if certifications:
                self._add_certifications_section(doc, certifications, generation_options, accent_rgb)

            # Add languages section
            languages = resume_content.get("languages", [])
            if languages:
                self._add_languages_section(doc, languages, generation_options, accent_rgb)

            # Add projects section
            projects = resume_content.get("projects", [])
            if projects:
                self._add_projects_section(doc, projects, generation_options, accent_rgb)

            # Save to bytes
            docx_bytes_io = io.BytesIO()
            doc.save(docx_bytes_io)
            docx_bytes = docx_bytes_io.getvalue()
            file_size = len(docx_bytes)

            # Estimate page count (rough approximation)
            page_count = max(1, file_size // 3000)  # ~3KB per page estimate

            logger.info(
                f"Successfully generated DOCX: {filename} "
                f"({file_size} bytes, ~{page_count} pages)"
            )

            return DOCXGenerationResult(
                success=True,
                docx_bytes=docx_bytes,
                filename=filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=file_size,
                page_count=page_count,
                metadata=metadata,
            )

        except Exception as e:
            logger.error(
                f"Failed to generate DOCX resume '{filename}': {e}",
                exc_info=True
            )
            return DOCXGenerationResult(
                success=False,
                filename=filename,
                error_message=f"DOCX generation failed: {str(e)}"
            )

    def _add_header_section(
        self,
        doc,
        personal_info: Dict[str, Any],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add header section with name and contact information.

        Args:
            doc: The Word document
            personal_info: Personal information dictionary
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        # Add name
        name = personal_info.get("full_name", "Resume")
        name_para = doc.add_paragraph()
        name_para.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = self._Pt(options.font_size_name)
        name_run.font.color.rgb = self._RGBColor(*accent_rgb)
        name_para.space_after = self._Pt(6)

        # Add title/headline if present
        title = personal_info.get("title")
        if title:
            title_para = doc.add_paragraph()
            title_para.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.size = self._Pt(12)
            title_run.italic = True
            title_para.space_after = self._Pt(8)

        # Add contact info on one line
        contact_parts = []
        if personal_info.get("email"):
            contact_parts.append(personal_info["email"])
        if personal_info.get("phone"):
            contact_parts.append(personal_info["phone"])
        if personal_info.get("location"):
            contact_parts.append(personal_info["location"])

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
            contact_text = " | ".join(contact_parts)
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = self._Pt(10)
            contact_para.space_after = self._Pt(4)

        # Add online profiles
        online_parts = []
        if personal_info.get("linkedin_url"):
            online_parts.append(f"LinkedIn: {personal_info['linkedin_url']}")
        if personal_info.get("github_url"):
            online_parts.append(f"GitHub: {personal_info['github_url']}")
        if personal_info.get("website_url"):
            online_parts.append(f"Website: {personal_info['website_url']}")

        if online_parts:
            online_para = doc.add_paragraph()
            online_para.alignment = self._WD_ALIGN_PARAGRAPH.CENTER
            online_text = " | ".join(online_parts)
            online_run = online_para.add_run(online_text)
            online_run.font.size = self._Pt(9)
            online_para.space_after = self._Pt(12)

    def _add_section_heading(
        self,
        doc,
        heading_text: str,
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add a section heading to the document.

        Args:
            doc: The Word document
            heading_text: Heading text
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        # Add horizontal line effect via border
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(heading_text.upper())
        heading_run.bold = True
        heading_run.font.size = self._Pt(options.font_size_heading)
        heading_run.font.color.rgb = self._RGBColor(*accent_rgb)

        # Add bottom border effect using underline
        heading_para.paragraph_format.space_before = self._Pt(12)
        heading_para.paragraph_format.space_after = self._Pt(6)

        # Add a separator line
        line_para = doc.add_paragraph()
        line_para.paragraph_format.space_before = self._Pt(0)
        line_para.paragraph_format.space_after = self._Pt(8)
        line_run = line_para.add_run("_" * 80)
        line_run.font.color.rgb = self._RGBColor(*accent_rgb)
        line_run.font.size = self._Pt(4)

    def _add_summary_section(
        self,
        doc,
        summary: str,
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add professional summary section.

        Args:
            doc: The Word document
            summary: Summary text
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Professional Summary", options, accent_rgb)

        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.size = self._Pt(options.font_size_body)
        summary_para.paragraph_format.space_after = self._Pt(8)

    def _add_work_experience_section(
        self,
        doc,
        work_experience: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add work experience section.

        Args:
            doc: The Word document
            work_experience: List of work experience entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Work Experience", options, accent_rgb)

        for exp in work_experience:
            # Job title and company line
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_after = self._Pt(0)

            title_run = header_para.add_run(exp.get("position", "Position"))
            title_run.bold = True
            title_run.font.size = self._Pt(options.font_size_body)

            company = exp.get("company")
            if company:
                company_run = header_para.add_run(f" at {company}")
                company_run.font.size = self._Pt(options.font_size_body)
                company_run.font.color.rgb = self._RGBColor(*accent_rgb)

            # Date and location line
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_before = self._Pt(0)
            date_para.paragraph_format.space_after = self._Pt(4)

            date_parts = []
            start_date = exp.get("start_date", "")
            end_date = exp.get("end_date", "Present") if not exp.get("is_current") else "Present"
            if start_date:
                date_parts.append(f"{start_date} - {end_date}")

            location = exp.get("location")
            if location:
                date_parts.append(location)

            if date_parts:
                date_run = date_para.add_run(" | ".join(date_parts))
                date_run.italic = True
                date_run.font.size = self._Pt(10)

            # Description
            description = exp.get("description")
            if description:
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.size = self._Pt(options.font_size_body - 1)
                desc_para.paragraph_format.space_after = self._Pt(4)

            # Highlights as bullet points
            highlights = exp.get("highlights", [])
            if highlights:
                for highlight in highlights:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet_para.add_run(highlight)
                    bullet_run.font.size = self._Pt(options.font_size_body - 1)
                    bullet_para.paragraph_format.space_after = self._Pt(2)

            # Add spacing after entry
            spacing_para = doc.add_paragraph()
            spacing_para.paragraph_format.space_after = self._Pt(6)

    def _add_education_section(
        self,
        doc,
        education: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add education section.

        Args:
            doc: The Word document
            education: List of education entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Education", options, accent_rgb)

        for edu in education:
            # Degree and institution line
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_after = self._Pt(0)

            degree = edu.get("degree", "Degree")
            field = edu.get("field_of_study")
            if field:
                degree_text = f"{degree} in {field}"
            else:
                degree_text = degree

            degree_run = header_para.add_run(degree_text)
            degree_run.bold = True
            degree_run.font.size = self._Pt(options.font_size_body)

            institution = edu.get("institution")
            if institution:
                inst_run = header_para.add_run(f" - {institution}")
                inst_run.font.size = self._Pt(options.font_size_body)
                inst_run.font.color.rgb = self._RGBColor(*accent_rgb)

            # Date and details line
            date_para = doc.add_paragraph()
            date_para.paragraph_format.space_before = self._Pt(0)
            date_para.paragraph_format.space_after = self._Pt(4)

            date_parts = []
            end_date = edu.get("end_date")
            if end_date:
                date_parts.append(end_date)

            gpa = edu.get("gpa")
            if gpa:
                date_parts.append(f"GPA: {gpa}")

            if date_parts:
                date_run = date_para.add_run(" | ".join(date_parts))
                date_run.italic = True
                date_run.font.size = self._Pt(10)

            # Honors
            honors = edu.get("honors", [])
            if honors:
                honors_para = doc.add_paragraph()
                honors_run = honors_para.add_run("Honors: " + ", ".join(honors))
                honors_run.font.size = self._Pt(options.font_size_body - 1)
                honors_run.italic = True
                honors_para.paragraph_format.space_after = self._Pt(4)

            # Description
            description = edu.get("description")
            if description:
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.size = self._Pt(options.font_size_body - 1)
                desc_para.paragraph_format.space_after = self._Pt(4)

            # Add spacing after entry
            spacing_para = doc.add_paragraph()
            spacing_para.paragraph_format.space_after = self._Pt(6)

    def _add_skills_section(
        self,
        doc,
        skills: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add skills section.

        Args:
            doc: The Word document
            skills: List of skill entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Skills", options, accent_rgb)

        # Group skills by category if available
        categories = {}
        for skill in skills:
            category = skill.get("category", "General")
            if category not in categories:
                categories[category] = []
            categories[category].append(skill)

        # If no categories, just list all skills
        if len(categories) == 1 and "General" in categories:
            skill_names = [s.get("name", str(s)) for s in skills]
            skills_para = doc.add_paragraph()
            skills_run = skills_para.add_run(", ".join(skill_names))
            skills_run.font.size = self._Pt(options.font_size_body)
            skills_para.paragraph_format.space_after = self._Pt(8)
        else:
            # Group by category
            for category, category_skills in categories.items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(f"{category}: ")
                cat_run.bold = True
                cat_run.font.size = self._Pt(options.font_size_body)

                skill_names = [s.get("name", str(s)) for s in category_skills]
                skills_run = cat_para.add_run(", ".join(skill_names))
                skills_run.font.size = self._Pt(options.font_size_body)
                cat_para.paragraph_format.space_after = self._Pt(4)

    def _add_certifications_section(
        self,
        doc,
        certifications: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add certifications section.

        Args:
            doc: The Word document
            certifications: List of certification entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Certifications", options, accent_rgb)

        for cert in certifications:
            cert_para = doc.add_paragraph()

            name_run = cert_para.add_run(cert.get("name", "Certification"))
            name_run.bold = True
            name_run.font.size = self._Pt(options.font_size_body)

            issuer = cert.get("issuer")
            if issuer:
                issuer_run = cert_para.add_run(f" - {issuer}")
                issuer_run.font.size = self._Pt(options.font_size_body)

            issue_date = cert.get("issue_date")
            if issue_date:
                date_run = cert_para.add_run(f" ({issue_date})")
                date_run.font.size = self._Pt(10)
                date_run.italic = True

            cert_para.paragraph_format.space_after = self._Pt(4)

    def _add_languages_section(
        self,
        doc,
        languages: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add languages section.

        Args:
            doc: The Word document
            languages: List of language entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Languages", options, accent_rgb)

        lang_parts = []
        for lang in languages:
            name = lang.get("name", str(lang))
            proficiency = lang.get("proficiency")
            if proficiency:
                lang_parts.append(f"{name} ({proficiency})")
            else:
                lang_parts.append(name)

        lang_para = doc.add_paragraph()
        lang_run = lang_para.add_run(", ".join(lang_parts))
        lang_run.font.size = self._Pt(options.font_size_body)
        lang_para.paragraph_format.space_after = self._Pt(8)

    def _add_projects_section(
        self,
        doc,
        projects: List[Dict[str, Any]],
        options: DOCXGenerationOptions,
        accent_rgb: tuple,
    ) -> None:
        """
        Add projects section.

        Args:
            doc: The Word document
            projects: List of project entries
            options: Generation options
            accent_rgb: Accent color as RGB tuple
        """
        self._add_section_heading(doc, "Projects", options, accent_rgb)

        for project in projects:
            # Project name and URL
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_after = self._Pt(0)

            name_run = header_para.add_run(project.get("name", "Project"))
            name_run.bold = True
            name_run.font.size = self._Pt(options.font_size_body)

            url = project.get("url")
            if url:
                url_run = header_para.add_run(f" - {url}")
                url_run.font.size = self._Pt(10)
                url_run.font.color.rgb = self._RGBColor(*accent_rgb)

            # Date
            start_date = project.get("start_date")
            end_date = project.get("end_date")
            if start_date:
                date_para = doc.add_paragraph()
                date_para.paragraph_format.space_before = self._Pt(0)
                date_para.paragraph_format.space_after = self._Pt(4)
                date_text = f"{start_date}"
                if end_date:
                    date_text += f" - {end_date}"
                date_run = date_para.add_run(date_text)
                date_run.italic = True
                date_run.font.size = self._Pt(10)

            # Description
            description = project.get("description")
            if description:
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.size = self._Pt(options.font_size_body - 1)
                desc_para.paragraph_format.space_after = self._Pt(4)

            # Technologies
            technologies = project.get("technologies", [])
            if technologies:
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run("Technologies: " + ", ".join(technologies))
                tech_run.font.size = self._Pt(options.font_size_body - 1)
                tech_run.italic = True
                tech_para.paragraph_format.space_after = self._Pt(4)

            # Add spacing after entry
            spacing_para = doc.add_paragraph()
            spacing_para.paragraph_format.space_after = self._Pt(6)

    async def generate_multiple_resumes(
        self,
        requests: List[Dict[str, Any]],
    ) -> List[DOCXGenerationResult]:
        """
        Generate multiple resume DOCX files in batch.

        Args:
            requests: List of dictionaries, each containing:
                - resume_content: Resume content dictionary
                - filename: Optional filename
                - candidate_name: Optional candidate name
                - metadata: Optional DOCX metadata
                - options: Optional DOCX generation options

        Returns:
            List of DOCXGenerationResult objects

        Example:
            >>> requests = [
            ...     {"resume_content": {...}, "candidate_name": "John Doe"},
            ...     {"resume_content": {...}, "candidate_name": "Jane Smith"},
            ... ]
            >>> results = await generator.generate_multiple_resumes(requests)
        """
        results = []

        for request in requests:
            result = await self.generate_resume_docx(
                resume_content=request.get("resume_content", {}),
                filename=request.get("filename"),
                candidate_name=request.get("candidate_name"),
                template_name=request.get("template_name"),
                metadata=request.get("metadata"),
                options=request.get("options"),
            )
            results.append(result)

        return results

    def health_check(self) -> Dict[str, Any]:
        """
        Check DOCX generator health and availability.

        Returns:
            Dictionary with health status and capabilities
        """
        return {
            "status": "healthy" if self.is_available() else "unavailable",
            "enabled": self.enabled,
            "docx_available": self.docx_available,
            "available_page_formats": list(self.PAGE_SIZES.keys()),
        }


def get_docx_generator() -> DOCXGenerator:
    """
    Get or create global DOCX generator service instance.

    Returns:
        Global DOCXGenerator instance

    Example:
        >>> from services.docx_generator import get_docx_generator
        >>> generator = get_docx_generator()
        >>> result = await generator.generate_resume_docx(resume_content, filename)
    """
    global _docx_generator
    if _docx_generator is None:
        _docx_generator = DOCXGenerator()
    return _docx_generator


async def generate_resume_docx(
    resume_content: Dict[str, Any],
    filename: Optional[str] = None,
    candidate_name: Optional[str] = None,
    template_name: Optional[str] = None,
    metadata: Optional[DOCXMetadata] = None,
    options: Optional[DOCXGenerationOptions] = None,
) -> DOCXGenerationResult:
    """
    Convenience function to generate a resume DOCX.

    This is a shortcut for using the DOCXGenerator class directly.

    Args:
        resume_content: Dictionary containing resume data
        filename: Optional filename for the DOCX
        candidate_name: Name of the candidate
        template_name: Name of the template used
        metadata: Optional DOCX metadata
        options: Optional DOCX generation options

    Returns:
        DOCXGenerationResult with generated DOCX or error details

    Example:
        >>> from services.docx_generator import generate_resume_docx
        >>> result = await generate_resume_docx(
        ...     resume_content={"personal_info": {...}, "work_experience": [...]},
        ...     candidate_name="John Doe"
        ... )
        >>> if result.success:
        ...     with open("resume.docx", "wb") as f:
        ...         f.write(result.docx_bytes)
    """
    generator = get_docx_generator()
    return await generator.generate_resume_docx(
        resume_content=resume_content,
        filename=filename,
        candidate_name=candidate_name,
        template_name=template_name,
        metadata=metadata,
        options=options,
    )
