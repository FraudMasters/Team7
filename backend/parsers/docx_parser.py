"""
DOCX text extraction from resume documents using python-docx.

This module provides functionality to extract text content from DOCX (Word) files,
with validation and error handling for resume processing. The parser handles
both text and table content, providing detailed error messages for malformed
or unsupported documents.

Includes complex layout detection for multi-column sections and tables,
providing structured information about the document's visual organization.
"""
import logging
import os
import re
from collections import Counter
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

logger = logging.getLogger(__name__)


@dataclass
class LayoutRegion:
    """
    Represents a detected region in a DOCX document.

    Attributes:
        region_type: Type of region (text, table, multi_column, header, footer)
        paragraph_index: Index of the starting paragraph
        confidence: Detection confidence (0.0 to 1.0)
        content: Extracted text content from the region
        row_count: Number of rows (for table type)
        column_count: Number of columns (for table type)
    """
    region_type: str
    paragraph_index: int = 0
    confidence: float = 1.0
    content: str = ""
    row_count: int = 0
    column_count: int = 0


@dataclass
class LayoutAnalysis:
    """
    Results of layout analysis for a DOCX document.

    Attributes:
        has_multi_column: Whether document has multi-column sections
        has_tables: Whether tables were detected
        regions: List of detected layout regions
        table_count: Total number of tables in document
        complexity_score: Overall layout complexity (0.0 to 1.0)
        warnings: List of warnings about layout issues
    """
    has_multi_column: bool = False
    has_tables: bool = False
    regions: List[LayoutRegion] = field(default_factory=list)
    table_count: int = 0
    complexity_score: float = 0.0
    warnings: List[str] = field(default_factory=list)


class DOCXParser:
    """
    Parser for extracting text from DOCX documents using python-docx.

    This class provides methods to extract raw text content from DOCX files,
    with validation for file integrity, structure validation, and text presence.
    It is designed to handle resume documents and extracts text from paragraphs
    and tables.

    Attributes:
        min_text_length: Minimum expected text length for valid resumes
        max_file_size_mb: Maximum allowed file size in megabytes
        extract_tables: Whether to extract text from tables
        extract_headers_footers: Whether to extract headers and footers

    Examples:
        >>> parser = DOCXParser()
        >>> result = parser.parse("/path/to/resume.docx")
        >>> print(result["text"])
        'John Doe\\nSoftware Engineer...'

        >>> # Parse from bytes
        >>> with open("resume.docx", "rb") as f:
        ...     docx_bytes = f.read()
        >>> result = parser.parse_bytes(docx_bytes)
        >>> print(result["text_length"])
        2543
    """

    def __init__(
        self,
        *,
        min_text_length: int = 50,
        max_file_size_mb: int = 10,
        extract_tables: bool = True,
        extract_headers_footers: bool = True,
        detect_layouts: bool = True,
    ):
        """
        Initialize the DOCX parser.

        Args:
            min_text_length: Minimum expected text length for valid resumes (default: 50 chars)
            max_file_size_mb: Maximum allowed file size in megabytes (default: 10MB)
            extract_tables: Whether to extract text from tables (default: True)
            extract_headers_footers: Whether to extract headers and footers (default: True)
            detect_layouts: Whether to detect complex layouts (default: True)
        """
        self.min_text_length = min_text_length
        self.max_file_size_mb = max_file_size_mb
        self.extract_tables = extract_tables
        self.extract_headers_footers = extract_headers_footers
        self.detect_layouts = detect_layouts

        # Check if python-docx is available
        try:
            from docx import Document  # type: ignore
            self.Document = Document
            logger.info("DOCXParser initialized with python-docx")
        except ImportError as e:
            raise ImportError(
                "python-docx is not installed. Install it with: pip install python-docx"
            ) from e

    def parse(
        self,
        file_path: Union[str, Path],
    ) -> Dict[str, Optional[Union[str, int, Dict[str, Union[str, int]]]]]:
        """
        Extract text from a DOCX file.

        This method reads a DOCX file from disk and extracts all text content
        from paragraphs and optionally from tables, headers, and footers.

        Args:
            file_path: Path to the DOCX file (string or Path object)

        Returns:
            Dictionary containing:
                - text: Extracted text content (None if failed)
                - text_length: Number of characters extracted (0 if failed)
                - paragraph_count: Number of paragraphs in document (0 if failed)
                - table_count: Number of tables extracted (0 if failed)
                - error: Error message if extraction failed (None if successful)
                - metadata: DOCX metadata dict (title, author, etc.)
                - layout_analysis: LayoutAnalysis object with detected regions (None if detection disabled)

        Raises:
            FileNotFoundError: If file_path does not exist
            ValueError: If file_path is not a DOCX file

        Examples:
            >>> parser = DOCXParser()
            >>> result = parser.parse("resume.docx")
            >>> if result["error"]:
            ...     print(f"Error: {result['error']}")
            ... else:
            ...     print(f"Extracted {result['text_length']} chars from {result['paragraph_count']} paragraphs")
        """
        # Validate file path
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"File not found: {file_path}",
                "metadata": None,
                "layout_analysis": None,
            }

        # Validate file extension
        if file_path.suffix.lower() not in (".docx", ".doc"):
            logger.error(f"Not a DOCX file: {file_path}")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"Not a DOCX file: {file_path.suffix}",
                "metadata": None,
                "layout_analysis": None,
            }

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            logger.error(f"File too large: {file_size_mb:.2f}MB (max: {self.max_file_size_mb}MB)")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"File too large: {file_size_mb:.2f}MB (max: {self.max_file_size_mb}MB)",
                "metadata": None,
                "layout_analysis": None,
            }

        try:
            logger.info(f"Parsing DOCX: {file_path} ({file_size_mb:.2f}MB)")

            # Read file bytes
            with open(file_path, "rb") as f:
                docx_bytes = f.read()

            # Parse using bytes method
            return self.parse_bytes(docx_bytes, filename=file_path.name)

        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"Parsing failed: {str(e)}",
                "metadata": None,
                "layout_analysis": None,
            }

    def parse_bytes(
        self,
        docx_bytes: bytes,
        *,
        filename: str = "document.docx",
    ) -> Dict[str, Optional[Union[str, int, Dict[str, Union[str, int]]]]]:
        """
        Extract text from DOCX bytes.

        This method extracts text from DOCX content provided as bytes, useful for
        handling file uploads or in-memory DOCX content.

        Args:
            docx_bytes: DOCX file content as bytes
            filename: Original filename (for logging and error messages)

        Returns:
            Dictionary containing:
                - text: Extracted text content (None if failed)
                - text_length: Number of characters extracted (0 if failed)
                - paragraph_count: Number of paragraphs in document (0 if failed)
                - table_count: Number of tables extracted (0 if failed)
                - error: Error message if extraction failed (None if successful)
                - metadata: DOCX metadata dict (title, author, etc.)
                - layout_analysis: LayoutAnalysis object with detected regions (None if detection disabled)

        Examples:
            >>> parser = DOCXParser()
            >>> with open("resume.docx", "rb") as f:
            ...     docx_bytes = f.read()
            >>> result = parser.parse_bytes(docx_bytes)
            >>> print(result["text"][:100])  # First 100 chars
        """
        # Validate input
        if not docx_bytes:
            logger.error("Empty DOCX bytes provided")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": "Empty DOCX bytes provided",
                "metadata": None,
                "layout_analysis": None,
            }

        # Check file size
        file_size_mb = len(docx_bytes) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            logger.error(f"DOCX too large: {file_size_mb:.2f}MB")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"DOCX too large: {file_size_mb:.2f}MB (max: {self.max_file_size_mb}MB)",
                "metadata": None,
                "layout_analysis": None,
            }

        try:
            # Create document from bytes
            docx_file = BytesIO(docx_bytes)
            doc = self.Document(docx_file)

            # Extract metadata
            metadata = self._extract_metadata(doc)

            # Extract text from paragraphs
            logger.info(f"Extracting text from {filename}")
            text_parts = []

            # Extract from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    paragraph_count += 1

            logger.debug(f"Extracted {paragraph_count} non-empty paragraphs")

            # Extract from tables if enabled
            table_count = 0
            if self.extract_tables:
                table_text = self._extract_table_text(doc)
                if table_text:
                    text_parts.append(table_text)
                    table_count = len(doc.tables)
                    logger.debug(f"Extracted text from {table_count} tables")

            # Extract from headers and footers if enabled
            if self.extract_headers_footers:
                header_footer_text = self._extract_header_footer_text(doc)
                if header_footer_text:
                    text_parts.append(header_footer_text)
                    logger.debug("Extracted headers and footers")

            # Combine all text
            full_text = "\n\n".join(text_parts)
            text_length = len(full_text)

            # Validate extracted text
            if text_length < self.min_text_length:
                logger.warning(f"Extracted text too short: {text_length} chars (min: {self.min_text_length})")

                if text_length == 0:
                    return {
                        "text": None,
                        "text_length": 0,
                        "paragraph_count": paragraph_count,
                        "table_count": table_count,
                        "error": "No extractable text found. Document may be empty or contain only images.",
                        "metadata": metadata,
                        "layout_analysis": None,
                    }
                else:
                    # Perform layout analysis even for short text
                    layout_analysis = None
                    if self.detect_layouts:
                        layout_analysis = self._analyze_layout(doc)

                    return {
                        "text": full_text,
                        "text_length": text_length,
                        "paragraph_count": paragraph_count,
                        "table_count": table_count,
                        "error": f"Extracted text too short: {text_length} chars (min: {self.min_text_length}). Document may be malformed.",
                        "metadata": metadata,
                        "layout_analysis": layout_analysis,
                    }

            logger.info(f"Successfully extracted {text_length} chars from {paragraph_count} paragraphs and {table_count} tables")

            # Perform layout analysis if enabled
            layout_analysis = None
            if self.detect_layouts:
                layout_analysis = self._analyze_layout(doc)
                if layout_analysis and layout_analysis.warnings:
                    for warning in layout_analysis.warnings:
                        logger.warning(f"Layout issue: {warning}")

            return {
                "text": full_text,
                "text_length": text_length,
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "error": None,
                "metadata": metadata,
                "layout_analysis": layout_analysis,
            }

        except Exception as e:
            logger.error(f"Failed to parse DOCX bytes from {filename}: {e}")
            return {
                "text": None,
                "text_length": 0,
                "paragraph_count": 0,
                "table_count": 0,
                "error": f"Parsing failed: {str(e)}",
                "metadata": None,
                "layout_analysis": None,
            }

    def _extract_metadata(self, doc) -> Dict[str, Optional[str]]:
        """
        Extract metadata from DOCX document.

        Args:
            doc: python-docx Document instance

        Returns:
            Dictionary with DOCX metadata fields
        """
        try:
            core_props = doc.core_properties

            return {
                "title": core_props.title or None,
                "author": core_props.author or None,
                "subject": core_props.subject or None,
                "keywords": core_props.keywords or None,
                "comments": core_props.comments or None,
                "category": core_props.category or None,
                "created": core_props.created or None,
                "modified": core_props.modified or None,
                "last_modified_by": core_props.last_modified_by or None,
                "revision": core_props.revision or None,
            }
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")

        return {
            "title": None,
            "author": None,
            "subject": None,
            "keywords": None,
            "comments": None,
            "category": None,
            "created": None,
            "modified": None,
            "last_modified_by": None,
            "revision": None,
        }

    def _extract_table_text(self, doc) -> Optional[str]:
        """
        Extract text from all tables in document.

        Args:
            doc: python-docx Document instance

        Returns:
            Combined text from all tables or None if no tables
        """
        if not doc.tables:
            return None

        table_texts = []
        for table_idx, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        cells_text.append(cell.text.strip())
                if cells_text:
                    rows_text.append(" | ".join(cells_text))

            if rows_text:
                table_text = "\n".join(rows_text)
                table_texts.append(f"[Table {table_idx + 1}]\n{table_text}")

        return "\n\n".join(table_texts) if table_texts else None

    def _extract_header_footer_text(self, doc) -> Optional[str]:
        """
        Extract text from headers and footers.

        Args:
            doc: python-docx Document instance

        Returns:
            Combined text from headers and footers or None if empty
        """
        text_parts = []

        try:
            # Extract from sections
            for section_idx, section in enumerate(doc.sections):
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Header Section {section_idx + 1}] {paragraph.text.strip()}")

                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"[Footer Section {section_idx + 1}] {paragraph.text.strip()}")

        except Exception as e:
            logger.warning(f"Failed to extract headers/footers: {e}")

        return "\n".join(text_parts) if text_parts else None

    def _analyze_layout(self, doc) -> LayoutAnalysis:
        """
        Analyze the layout of the DOCX document.

        Detects tables and analyzes their structure, identifies potential
        multi-column sections, and calculates complexity score.

        Args:
            doc: python-docx Document instance

        Returns:
            LayoutAnalysis object with detected regions and complexity score
        """
        analysis = LayoutAnalysis()
        regions = []

        try:
            # Analyze tables
            table_regions = self._detect_tables_detailed(doc)
            if table_regions:
                analysis.has_tables = True
                analysis.table_count = len(table_regions)
                regions.extend(table_regions)

            # Analyze potential multi-column sections
            column_regions = self._detect_columns_detailed(doc)
            if column_regions:
                analysis.has_multi_column = True
                regions.extend(column_regions)

            analysis.regions = regions
            analysis.complexity_score = self._calculate_complexity(analysis)

            # Add warnings for complex layouts
            if analysis.has_multi_column:
                analysis.warnings.append("Multi-column sections detected. Text flow may be affected.")

            if analysis.table_count > 3:
                analysis.warnings.append(f"Many tables detected ({analysis.table_count}). Consider verifying extracted data.")

        except Exception as e:
            logger.warning(f"Failed to analyze layout: {e}")
            analysis.warnings.append(f"Layout analysis failed: {str(e)}")

        return analysis

    def _detect_tables_detailed(self, doc) -> List[LayoutRegion]:
        """
        Detect and analyze tables in detail.

        Args:
            doc: python-docx Document instance

        Returns:
            List of LayoutRegion objects for detected tables
        """
        regions = []

        if not doc.tables:
            return regions

        for table_idx, table in enumerate(doc.tables):
            try:
                # Count rows and columns
                row_count = len(table.rows)
                col_count = len(table.columns) if table.rows else 0

                # Extract table content for analysis
                cells_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        cells_content.append(" | ".join(row_text))

                table_text = "\n".join(cells_content)

                # Determine table position (approximate by paragraph index)
                # This is a heuristic - we look for table references in paragraphs
                para_index = self._estimate_table_position(doc, table_idx)

                # Calculate confidence based on table structure
                confidence = 1.0 if row_count > 1 and col_count > 1 else 0.6

                region = LayoutRegion(
                    region_type="table",
                    paragraph_index=para_index,
                    confidence=confidence,
                    content=table_text,
                    row_count=row_count,
                    column_count=col_count,
                )
                regions.append(region)
                logger.debug(f"Detected table {table_idx + 1}: {row_count} rows x {col_count} cols")

            except Exception as e:
                logger.warning(f"Failed to analyze table {table_idx}: {e}")

        return regions

    def _estimate_table_position(self, doc, table_idx: int) -> int:
        """
        Estimate the paragraph index where a table appears.

        This is a heuristic approach since python-docx doesn't provide
        direct table position in document flow.

        Args:
            doc: python-docx Document instance
            table_idx: Index of the table in doc.tables

        Returns:
            Estimated paragraph index
        """
        # Simple heuristic: assume tables are roughly evenly distributed
        # through the document
        total_paras = len(doc.paragraphs)
        total_tables = len(doc.tables)

        if total_tables == 0:
            return 0

        # Estimate position based on table index
        estimated_position = int((table_idx + 1) / total_tables * total_paras)
        return min(estimated_position, total_paras - 1)

    def _detect_columns_detailed(self, doc) -> List[LayoutRegion]:
        """
        Detect potential multi-column sections in the document.

        Uses heuristics based on paragraph formatting and content patterns.
        Note: python-docx has limited support for column detection, so this
        is a best-effort heuristic.

        Args:
            doc: python-docx Document instance

        Returns:
            List of LayoutRegion objects for detected multi-column sections
        """
        regions = []

        # Check sections for column settings
        try:
            for section_idx, section in enumerate(doc.sections):
                # Get column count from section properties
                # Note: This only works if columns are set at section level
                # In many cases, columns are set per-paragraph which is harder to detect
                column_count = 1
                try:
                    # Try to access column properties if available
                    if hasattr(section, '_sectPr'):
                        cols = section._sectPr.find(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols'
                        )
                        if cols is not None:
                            num_cols = cols.get('num')
                            if num_cols:
                                column_count = int(num_cols)
                except Exception:
                    pass

                if column_count > 1:
                    # Find paragraphs in this section
                    start_para = 0
                    if section_idx > 0:
                        # Estimate start based on previous sections
                        start_para = len(doc.paragraphs) // len(doc.sections) * section_idx
                    end_para = len(doc.paragraphs) // len(doc.sections) * (section_idx + 1)

                    section_text = "\n".join(
                        p.text for p in doc.paragraphs[start_para:end_para]
                        if p.text.strip()
                    )

                    region = LayoutRegion(
                        region_type="multi_column",
                        paragraph_index=start_para,
                        confidence=0.8,  # Lower confidence since detection is heuristic
                        content=section_text,
                        column_count=column_count,
                    )
                    regions.append(region)
                    logger.debug(f"Detected {column_count}-column section at paragraph {start_para}")

        except Exception as e:
            logger.warning(f"Failed to detect columns: {e}")

        # Additional heuristic: Look for tabular patterns in paragraph text
        # This can indicate content that was formatted as columns
        tabular_regions = self._detect_tabular_paragraphs(doc)
        regions.extend(tabular_regions)

        return regions

    def _detect_tabular_paragraphs(self, doc) -> List[LayoutRegion]:
        """
        Detect paragraphs that contain tabular/columnar data.

        Args:
            doc: python-docx Document instance

        Returns:
            List of LayoutRegion objects for detected tabular content
        """
        regions = []
        tabular_pattern = re.compile(r'(\S+\s{2,}\S+)|(\S+\t+\S+)')

        current_group = []
        group_start = 0

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph looks like tabular data
            # (multiple consecutive spaces or tabs separating content)
            if tabular_pattern.search(text):
                if not current_group:
                    group_start = idx
                current_group.append((idx, text))
            else:
                if len(current_group) >= 3:
                    # Found a group of tabular paragraphs
                    group_text = "\n".join(t for _, t in current_group)

                    # Estimate column count by analyzing spacing
                    col_count = self._estimate_column_count([t for _, t in current_group])

                    region = LayoutRegion(
                        region_type="multi_column",
                        paragraph_index=group_start,
                        confidence=0.6,
                        content=group_text,
                        column_count=col_count,
                    )
                    regions.append(region)
                    logger.debug(f"Detected tabular content at paragraph {group_start}")

                current_group = []

        # Check final group
        if len(current_group) >= 3:
            group_text = "\n".join(t for _, t in current_group)
            col_count = self._estimate_column_count([t for _, t in current_group])

            region = LayoutRegion(
                region_type="multi_column",
                paragraph_index=group_start,
                confidence=0.6,
                content=group_text,
                column_count=col_count,
            )
            regions.append(region)

        return regions

    def _estimate_column_count(self, texts: List[str]) -> int:
        """
        Estimate the number of columns in tabular text.

        Args:
            texts: List of text strings from tabular paragraphs

        Returns:
            Estimated column count
        """
        if not texts:
            return 1

        # Count separators (multiple spaces or tabs) in each line
        separator_counts = []
        for text in texts:
            # Count tab separators
            tab_count = text.count('\t')
            # Count multi-space separators (3+ spaces)
            multi_space_count = len(re.findall(r'\s{3,}', text))
            separator_counts.append(max(tab_count, multi_space_count))

        # Most common separator count + 1 = column count
        if separator_counts:
            common = Counter(separator_counts).most_common(1)
            if common:
                return common[0][0] + 1

        return 1

    def _calculate_complexity(self, analysis: LayoutAnalysis) -> float:
        """
        Calculate overall layout complexity score.

        Args:
            analysis: LayoutAnalysis object with detected regions

        Returns:
            Complexity score from 0.0 (simple) to 1.0 (very complex)
        """
        score = 0.0

        # Multi-column adds complexity
        if analysis.has_multi_column:
            score += 0.25

        # Tables add complexity
        if analysis.has_tables:
            score += min(0.35, analysis.table_count * 0.1)

        # Large number of regions indicates complexity
        if len(analysis.regions) > 5:
            score += 0.2

        # Many tables indicate complex document
        if analysis.table_count > 5:
            score += 0.2

        return min(1.0, score)

    def validate_docx(
        self,
        file_path: Union[str, Path],
    ) -> Dict[str, Union[bool, str, Optional[str]]]:
        """
        Validate DOCX file before parsing.

        This method performs quick validation checks without extracting text,
        useful for pre-flight validation.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing:
                - is_valid: Whether DOCX passes validation
                - error: Error message if invalid (None if valid)
                - file_size_mb: File size in MB
                - reason: Human-readable validation result

        Examples:
            >>> parser = DOCXParser()
            >>> result = parser.validate_docx("resume.docx")
            >>> if result["is_valid"]:
            ...     print("DOCX is valid for parsing")
            ... else:
            ...     print(f"Invalid: {result['error']}")
        """
        file_path = Path(file_path)

        # Check existence
        if not file_path.exists():
            return {
                "is_valid": False,
                "error": "File not found",
                "file_size_mb": 0,
                "reason": f"File does not exist: {file_path}",
            }

        # Check extension
        if file_path.suffix.lower() not in (".docx", ".doc"):
            return {
                "is_valid": False,
                "error": "Invalid file type",
                "file_size_mb": file_path.stat().st_size / (1024 * 1024),
                "reason": f"Not a DOCX file: {file_path.suffix}",
            }

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            return {
                "is_valid": False,
                "error": "File too large",
                "file_size_mb": file_size_mb,
                "reason": f"File size {file_size_mb:.2f}MB exceeds maximum {self.max_file_size_mb}MB",
            }

        # Check if file is empty
        if file_size_mb == 0:
            return {
                "is_valid": False,
                "error": "Empty file",
                "file_size_mb": 0,
                "reason": "DOCX file is empty (0 bytes)",
            }

        # Try to open DOCX with python-docx
        try:
            with open(file_path, "rb") as f:
                doc = self.Document(f)

                # Check if has paragraphs
                if len(doc.paragraphs) == 0:
                    return {
                        "is_valid": False,
                        "error": "No content",
                        "file_size_mb": file_size_mb,
                        "reason": "DOCX contains no paragraphs",
                    }

        except Exception as e:
            return {
                "is_valid": False,
                "error": "Invalid DOCX",
                "file_size_mb": file_size_mb,
                "reason": f"Failed to open DOCX: {str(e)}",
            }

        return {
            "is_valid": True,
            "error": None,
            "file_size_mb": file_size_mb,
            "reason": f"Valid DOCX with {file_size_mb:.2f}MB",
        }
